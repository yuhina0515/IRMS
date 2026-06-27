import docx
import sys

def dump_docx(filepath):
    doc = docx.Document(filepath)
    print("=== Paragraphs ===")
    for i, p in enumerate(doc.paragraphs):
        print(f"P[{i}]: {p.text}")
    
    print("\n=== Tables ===")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table[{t_idx}]:")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.replace('\n', ' ') for cell in row.cells]
            print(f"  Row[{r_idx}]: {cells}")

if __name__ == '__main__':
    dump_docx(sys.argv[1])
