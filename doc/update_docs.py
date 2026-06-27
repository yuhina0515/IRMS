import docx
import os

docs_info = {
    "5.docx": {
        "title": "分工一：邊緣端與韌體開發工程師 (Edge Firmware Developer)",
        "content": [
            "主要負責 ESP32 核心韌體開發、測試與維護。",
            "處理 MPU6050 雙慣性感測器 (IMU) 的資料擷取、I2C 總線狀態監控及防錯復原機制。",
            "實作與優化互補濾波演算法 (Complementary Filter)，處理動態校準與零點偏移，確保角度計算精準。",
            "管理 FreeRTOS 多執行緒任務的排程與資源分配 (包含 Task_Sensor, Task_Logic, Task_Comm, Task_LED)。",
            "實作 BLE 藍牙通訊，將 6 軸角度數據推播至應用端，並接收參數設定寫入 NVS 持久化儲存。"
        ]
    },
    "37.docx": {
        "title": "分工二：桌面端與後端應用開發工程師 (Desktop & Backend Application Developer)",
        "content": [
            "主要負責 Electron 桌面應用程式後端邏輯與系統架構設計。",
            "實作 Express 本地伺服器與 REST API，作為前端與資料庫溝通的橋樑。",
            "管理 SQLite 本地資料庫 (sessions 與 sensor_data 資料表)，設計關聯結構以完整保存復健歷程與高頻感測數據。",
            "開發藍牙自動配對邏輯，監聽連線狀態與異常處理 (如 I2C 斷線時的錯誤攔截與復原處理)。",
            "負責實作歷史數據的 CSV 導出邏輯，確保匯出數據完整性與格式正確性。"
        ]
    },
    "34.docx": {
        "title": "分工三：前端介面與資料視覺化工程師 (Frontend & Data Visualization Developer)",
        "content": [
            "主要負責 Electron App 的前端使用者介面 (UI) 設計及使用者體驗 (UX) 優化。",
            "使用 Vanilla JS 與 Web Bluetooth API 介接藍牙資料，處理即時數據流解碼與呈現。",
            "整合 Chart.js 開發高效能即時關節角度折線圖，並實作 Varus/Valgus 內外翻角度等即時視覺化卡片。",
            "實作多協定指定動作 (Designated Action) 的邏輯判定介面與參數同步設定面板。",
            "套用深色模式與玻璃擬物化 (Glassmorphism) 風格，提升系統的整體視覺質感與互動回饋（如紅/綠色遮罩警告燈號）。"
        ]
    }
}

base_path = r"c:\Users\Yuhina\Documents\IRMS\doc"

for filename, info in docs_info.items():
    filepath = os.path.join(base_path, filename)
    doc = docx.Document()
    doc.add_heading(info["title"], level=1)
    doc.add_paragraph("根據 IRMS (智慧復健監測系統) 專案的軟硬體架構，您的主要負責項目如下：\n")
    for item in info["content"]:
        doc.add_paragraph(item, style='List Bullet')
    doc.save(filepath)
    print(f"Successfully updated {filepath}")
